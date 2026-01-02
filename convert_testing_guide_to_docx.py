"""
Convert LOCAL_TESTING_GUIDE.md to a professionally formatted DOCX file.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_shading_to_paragraph(paragraph, color='F0F0F0'):
    """Add background shading to a paragraph."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    paragraph._element.get_or_add_pPr().append(shading_elm)

def create_local_testing_guide_docx():
    """Create a formatted DOCX version of the local testing guide."""
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Local Testing Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Heart Disease MLOps Project - Complete Setup and Access Instructions')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.italic = True
    
    doc.add_paragraph()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('1. Prerequisites', style='List Number')
    doc.add_paragraph('2. Initial Setup', style='List Number')
    doc.add_paragraph('3. Running the ML Pipeline', style='List Number')
    doc.add_paragraph('4. Kubernetes Deployment', style='List Number')
    doc.add_paragraph('5. Accessing All Services', style='List Number')
    doc.add_paragraph('6. Testing the Application', style='List Number')
    doc.add_paragraph('7. Troubleshooting', style='List Number')
    
    doc.add_page_break()
    
    # Section 1: Prerequisites
    doc.add_heading('1. Prerequisites', level=1)
    doc.add_paragraph('Before starting, ensure you have the following installed:')
    doc.add_paragraph('Python 3.9+ - https://www.python.org/downloads/', style='List Bullet')
    doc.add_paragraph('Docker Desktop - https://www.docker.com/products/docker-desktop/', style='List Bullet')
    doc.add_paragraph('Kubernetes - Enable in Docker Desktop settings', style='List Bullet')
    doc.add_paragraph('Git - https://git-scm.com/downloads', style='List Bullet')
    
    doc.add_heading('Verify Installations', level=2)
    code = doc.add_paragraph()
    run = code.add_run('python --version\ndocker --version\nkubectl version --client\ngit --version')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.5)
    add_shading_to_paragraph(code)
    
    # Section 2: Initial Setup
    doc.add_heading('2. Initial Setup', level=1)
    
    doc.add_heading('Step 1: Clone the Repository', level=2)
    code = doc.add_paragraph()
    run = code.add_run('git clone https://github.com/ssrikantasahoo/heart-disease-mlops-final.git\ncd heart-disease-mlops-final')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.5)
    add_shading_to_paragraph(code)
    
    doc.add_heading('Step 2: Create Virtual Environment', level=2)
    doc.add_paragraph('Windows:')
    code = doc.add_paragraph()
    run = code.add_run('python -m venv venv\n.\\venv\\Scripts\\activate')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.5)
    add_shading_to_paragraph(code)
    
    doc.add_paragraph('Mac/Linux:')
    code = doc.add_paragraph()
    run = code.add_run('python -m venv venv\nsource venv/bin/activate')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.5)
    add_shading_to_paragraph(code)
    
    doc.add_heading('Step 3: Install Dependencies', level=2)
    code = doc.add_paragraph()
    run = code.add_run('pip install -r requirements.txt')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.5)
    add_shading_to_paragraph(code)
    
    # Section 3: Running ML Pipeline
    doc.add_heading('3. Running the ML Pipeline', level=1)
    doc.add_paragraph('The ML pipeline performs data acquisition, preprocessing, model training, and experiment tracking.')
    
    doc.add_heading('Step 1: Run the Complete Pipeline', level=2)
    code = doc.add_paragraph()
    run = code.add_run('python run_local_pipeline.py')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.5)
    add_shading_to_paragraph(code)
    
    doc.add_paragraph('This script will:')
    doc.add_paragraph('Download the UCI Heart Disease dataset', style='List Bullet')
    doc.add_paragraph('Preprocess the data (handle missing values, encode features)', style='List Bullet')
    doc.add_paragraph('Train models (Logistic Regression and Random Forest)', style='List Bullet')
    doc.add_paragraph('Log experiments to MLflow', style='List Bullet')
    doc.add_paragraph('Package the best model', style='List Bullet')
    doc.add_paragraph('Run unit tests', style='List Bullet')
    
    doc.add_heading('Step 2: View MLflow Experiments', level=2)
    code = doc.add_paragraph()
    run = code.add_run('mlflow ui')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.5)
    add_shading_to_paragraph(code)
    
    para = doc.add_paragraph()
    para.add_run('Access MLflow: ').bold = True
    para.add_run('http://localhost:5000')
    
    # Section 4: Kubernetes Deployment
    doc.add_heading('4. Kubernetes Deployment', level=1)
    
    doc.add_heading('Step 1: Build Docker Images', level=2)
    code = doc.add_paragraph()
    run = code.add_run('docker build -t heart-api .\ndocker build -f Dockerfile.ui -t heart-ui .')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.5)
    add_shading_to_paragraph(code)
    
    doc.add_heading('Step 2: Deploy to Kubernetes', level=2)
    code = doc.add_paragraph()
    run = code.add_run('kubectl apply -f k8s/')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.5)
    add_shading_to_paragraph(code)
    
    doc.add_heading('Step 3: Verify Deployments', level=2)
    code = doc.add_paragraph()
    run = code.add_run('kubectl get pods\nkubectl get services')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.5)
    add_shading_to_paragraph(code)
    
    doc.add_paragraph('Wait for all pods to show Running status (1-2 minutes).')
    
    doc.add_page_break()
    
    # Section 5: Accessing All Services
    doc.add_heading('5. Accessing All Services', level=1)
    
    # Service 1: Web UI
    doc.add_heading('5.1 Prediction Web UI', level=2)
    para = doc.add_paragraph()
    para.add_run('URL: ').bold = True
    para.add_run('http://localhost:30081')
    
    para = doc.add_paragraph()
    para.add_run('Purpose: ').bold = True
    para.add_run('Interactive form to input patient data and get predictions')
    
    doc.add_paragraph('How to use:')
    doc.add_paragraph('Open http://localhost:30081 in your browser', style='List Number')
    doc.add_paragraph('Fill in patient health metrics', style='List Number')
    doc.add_paragraph('Click "Predict" button', style='List Number')
    doc.add_paragraph('View prediction result', style='List Number')
    
    # Service 2: API Docs
    doc.add_heading('5.2 FastAPI Documentation', level=2)
    para = doc.add_paragraph()
    para.add_run('URL: ').bold = True
    para.add_run('http://localhost:30080/docs')
    
    para = doc.add_paragraph()
    para.add_run('Purpose: ').bold = True
    para.add_run('Interactive API documentation with Swagger UI')
    
    doc.add_paragraph('How to use:')
    doc.add_paragraph('Open http://localhost:30080/docs', style='List Number')
    doc.add_paragraph('Expand /predict endpoint', style='List Number')
    doc.add_paragraph('Click "Try it out"', style='List Number')
    doc.add_paragraph('Enter sample patient data', style='List Number')
    doc.add_paragraph('Click "Execute"', style='List Number')
    doc.add_paragraph('View response', style='List Number')
    
    # Service 3: MLflow
    doc.add_heading('5.3 MLflow UI', level=2)
    para = doc.add_paragraph()
    para.add_run('URL: ').bold = True
    para.add_run('http://localhost:30050')
    
    para = doc.add_paragraph()
    para.add_run('Purpose: ').bold = True
    para.add_run('Experiment tracking and model registry')
    
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('Compare multiple runs side-by-side', style='List Bullet')
    doc.add_paragraph('View hyperparameters and metrics', style='List Bullet')
    doc.add_paragraph('Download model artifacts', style='List Bullet')
    doc.add_paragraph('Track model versions', style='List Bullet')
    
    # Service 4: Prometheus
    doc.add_heading('5.4 Prometheus Metrics', level=2)
    para = doc.add_paragraph()
    para.add_run('URL: ').bold = True
    para.add_run('http://localhost:30090')
    
    para = doc.add_paragraph()
    para.add_run('Purpose: ').bold = True
    para.add_run('Raw metrics collection and monitoring')
    
    doc.add_paragraph('Available Metrics:')
    doc.add_paragraph('api_requests_total - Total API requests', style='List Bullet')
    doc.add_paragraph('up{job="heart-api"} - API health status', style='List Bullet')
    
    # Service 5: Grafana
    doc.add_heading('5.5 Grafana Dashboard', level=2)
    para = doc.add_paragraph()
    para.add_run('URL: ').bold = True
    para.add_run('http://localhost:30300')
    
    para = doc.add_paragraph()
    para.add_run('Credentials: ').bold = True
    para.add_run('Username: admin, Password: admin')
    
    para = doc.add_paragraph()
    para.add_run('Purpose: ').bold = True
    para.add_run('Rich visualization dashboards with real-time updates')
    
    doc.add_paragraph('How to use:')
    doc.add_paragraph('Navigate to http://localhost:30300', style='List Number')
    doc.add_paragraph('Login with admin/admin', style='List Number')
    doc.add_paragraph('Click Dashboards → Browse', style='List Number')
    doc.add_paragraph('Select "Heart Disease API Metrics"', style='List Number')
    
    doc.add_paragraph('Dashboard Panels:')
    doc.add_paragraph('API Request Rate - Requests per second', style='List Bullet')
    doc.add_paragraph('Total API Requests - Cumulative count (gauge)', style='List Bullet')
    doc.add_paragraph('API Health Status - Green (Up) or Red (Down)', style='List Bullet')
    doc.add_paragraph('Cumulative Requests Over Time - Historical trend', style='List Bullet')
    
    doc.add_page_break()
    
    # Section 6: Testing
    doc.add_heading('6. Testing the Application', level=1)
    
    doc.add_heading('Test 1: Make a Prediction via Web UI', level=2)
    doc.add_paragraph('Open http://localhost:30081', style='List Number')
    doc.add_paragraph('Enter patient data (age: 63, sex: 1, cp: 3, etc.)', style='List Number')
    doc.add_paragraph('Click "Predict"', style='List Number')
    doc.add_paragraph('Verify you get a prediction result', style='List Number')
    
    doc.add_heading('Test 2: Make a Prediction via API', level=2)
    doc.add_paragraph('Using PowerShell:')
    code = doc.add_paragraph()
    run = code.add_run('Invoke-WebRequest -Uri "http://localhost:30080/predict" `\n  -Method POST `\n  -ContentType "application/json" `\n  -Body \'{"age": 63, "sex": 1, "cp": 3, "trestbps": 145, "chol": 233, "fbs": 1, "restecg": 0, "thalach": 150, "exang": 0, "oldpeak": 2.3, "slope": 0, "ca": 0, "thal": 1}\'')
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    code.paragraph_format.left_indent = Inches(0.5)
    add_shading_to_paragraph(code)
    
    doc.add_heading('Test 3: Verify Metrics in Grafana', level=2)
    doc.add_paragraph('Make 10-20 predictions', style='List Number')
    doc.add_paragraph('Open Grafana: http://localhost:30300', style='List Number')
    doc.add_paragraph('Navigate to "Heart Disease API Metrics" dashboard', style='List Number')
    doc.add_paragraph('Verify metrics are updating', style='List Number')
    
    # Section 7: Troubleshooting
    doc.add_heading('7. Troubleshooting', level=1)
    
    doc.add_heading('Issue 1: Pods Not Starting', level=2)
    code = doc.add_paragraph()
    run = code.add_run('kubectl get pods\nkubectl logs <pod-name>\nkubectl describe pod <pod-name>')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.5)
    add_shading_to_paragraph(code)
    
    doc.add_paragraph('Solution: Wait 1-2 minutes, check Docker Desktop is running')
    
    doc.add_heading('Issue 2: Cannot Access Services', level=2)
    code = doc.add_paragraph()
    run = code.add_run('kubectl get services')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.5)
    add_shading_to_paragraph(code)
    
    doc.add_paragraph('Solution: Restart Docker Desktop or redeploy')
    
    doc.add_heading('Issue 3: Grafana Shows "No Data"', level=2)
    doc.add_paragraph('Solution:')
    doc.add_paragraph('Make some predictions to generate traffic', style='List Number')
    doc.add_paragraph('Check Prometheus targets at http://localhost:30090/targets', style='List Number')
    doc.add_paragraph('Wait 10-15 seconds for metrics to appear', style='List Number')
    
    # Quick Reference Table
    doc.add_page_break()
    doc.add_heading('Quick Reference - All Access URLs', level=1)
    
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Service'
    header_cells[1].text = 'URL'
    header_cells[2].text = 'Credentials'
    header_cells[3].text = 'Purpose'
    
    # Data rows
    services = [
        ('Web UI', 'http://localhost:30081', 'None', 'Make predictions'),
        ('API Docs', 'http://localhost:30080/docs', 'None', 'API documentation'),
        ('MLflow', 'http://localhost:30050', 'None', 'Experiment tracking'),
        ('Prometheus', 'http://localhost:30090', 'None', 'Metrics collection'),
        ('Grafana', 'http://localhost:30300', 'admin/admin', 'Monitoring dashboards'),
        ('API (ReDoc)', 'http://localhost:30080/redoc', 'None', 'Alternative API docs')
    ]
    
    for i, (service, url, creds, purpose) in enumerate(services, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = service
        row_cells[1].text = url
        row_cells[2].text = creds
        row_cells[3].text = purpose
    
    # Save document
    output_path = 'docs/LOCAL_TESTING_GUIDE.docx'
    doc.save(output_path)
    
    print(f"✅ DOCX guide created successfully: {output_path}")
    print(f"📄 Document contains comprehensive testing instructions")
    print(f"📊 Includes all service access URLs and credentials")
    
    return output_path

if __name__ == "__main__":
    create_local_testing_guide_docx()
